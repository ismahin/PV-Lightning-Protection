"""Build the step-by-step Simscape results report from measured assets."""

from __future__ import annotations

import csv
import json
from datetime import date
from pathlib import Path

import matplotlib.pyplot as plt
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "docs" / "report_assets"
OUTPUT = ROOT / "docs" / "PV_Simscape_Step_by_Step_Report.docx"


def load_data():
    with (ASSETS / "measured_metrics.csv").open(newline="", encoding="utf-8-sig") as handle:
        metrics = {row["metric"]: float(row["value"]) for row in csv.DictReader(handle)}
    with (ASSETS / "summary.json").open(encoding="utf-8") as handle:
        summary = json.load(handle)
    return metrics, summary


def shade(cell, fill):
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def set_cell_text(cell, text, bold=False, color=None, size=9):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_caption(document, text):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.bold = True
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(70, 70, 70)


def add_picture(document, filename, caption, width=6.85):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(ASSETS / filename), width=Inches(width))
    add_caption(document, caption)


def add_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.add_run(text)


def add_step_heading(document, number, title):
    heading = document.add_heading(f"Step {number} — {title}", level=1)
    heading.paragraph_format.space_before = Pt(4)
    heading.paragraph_format.space_after = Pt(6)


def create_flow_diagram(metrics, summary):
    output = ASSETS / "00_system_flow.png"
    fig, ax = plt.subplots(figsize=(14, 6), dpi=180)
    ax.set_xlim(0, 14)
    ax.set_ylim(0, 6)
    ax.axis("off")
    fig.patch.set_facecolor("white")

    boxes = [
        (0.2, 3.2, 1.7, 1.25, "PV array", f"{metrics['PV voltage before surge']:.2f} V\n{metrics['PV current before surge']:.2f} A"),
        (2.2, 3.2, 1.7, 1.25, "MPPT + boost", f"Tracks maximum power\n≈ {metrics['PV power delivered to MPPT/boost']:.1f} W"),
        (4.2, 3.2, 1.8, 1.25, "Injection node", f"Surge peak\n{metrics['Injection-node voltage peak']:.1f} V"),
        (6.3, 3.2, 1.65, 1.25, "SPD1 stage", f"Knee {metrics['SPD1 conduction threshold']:.0f} V\n{metrics['SPD1 diverted-current peak']:.0f} A to ground"),
        (8.25, 3.2, 1.65, 1.25, "SPD2 stage", f"Knee {metrics['SPD2 conduction threshold']:.0f} V\n{metrics['SPD2 diverted-current peak']:.1f} A to ground"),
        (10.2, 3.2, 1.65, 1.25, "Protected bus", f"Normal {metrics['Protected DC bus before surge']:.2f} V\nPeak {metrics['Protected DC-bus peak']:.2f} V"),
        (12.15, 3.2, 1.65, 1.25, "Inverter + load", f"{metrics['AC output voltage RMS before surge']:.1f} V RMS\n{metrics['AC output current RMS before surge']:.2f} A RMS"),
    ]
    colors = ["#DFF3FF", "#DDF7E3", "#FFE9C6", "#FFE0E0", "#FFE8D5", "#E5E8FF", "#E5F4EC"]

    for index, (x, y, width, height, title, detail) in enumerate(boxes):
        patch = FancyBboxPatch((x, y), width, height, boxstyle="round,pad=0.04,rounding_size=0.08",
                               linewidth=1.5, edgecolor="#334155", facecolor=colors[index])
        ax.add_patch(patch)
        ax.text(x + width / 2, y + 0.82, title, ha="center", va="center", fontsize=11, weight="bold")
        ax.text(x + width / 2, y + 0.34, detail, ha="center", va="center", fontsize=9)
        if index < len(boxes) - 1:
            next_x = boxes[index + 1][0]
            ax.add_patch(FancyArrowPatch((x + width, y + height / 2), (next_x, y + height / 2),
                                         arrowstyle="-|>", mutation_scale=15, linewidth=1.5, color="#334155"))

    ax.add_patch(FancyBboxPatch((4.22, 5.05), 1.75, 0.65, boxstyle="round,pad=0.04",
                                linewidth=1.5, edgecolor="#9A3412", facecolor="#FFF1E6"))
    ax.text(5.095, 5.38, f"Lightning source: 10 kA, {summary['frontTime_us']:.0f}/{summary['halfValueTime_us']:.0f} µs",
            ha="center", va="center", fontsize=9.5, weight="bold", color="#9A3412")
    ax.add_patch(FancyArrowPatch((5.1, 5.05), (5.1, 4.48), arrowstyle="-|>", mutation_scale=16,
                                 linewidth=1.8, color="#C2410C"))

    for x, label in [(7.12, "SPD1 surge\nto ground"), (9.08, "SPD2 residual\nto ground")]:
        ax.add_patch(FancyArrowPatch((x, 3.18), (x, 1.65), arrowstyle="-|>", mutation_scale=16,
                                     linewidth=1.8, color="#B91C1C"))
        ax.text(x, 1.34, label, ha="center", va="center", fontsize=9, color="#991B1B", weight="bold")
        ax.plot([x - 0.27, x + 0.27], [0.95, 0.95], color="#334155", linewidth=1.5)
        ax.plot([x - 0.18, x + 0.18], [0.78, 0.78], color="#334155", linewidth=1.5)
        ax.plot([x - 0.09, x + 0.09], [0.61, 0.61], color="#334155", linewidth=1.5)

    ax.add_patch(FancyBboxPatch((10.2, 1.0), 1.65, 0.9, boxstyle="round,pad=0.04",
                                linewidth=1.5, edgecolor="#0F766E", facecolor="#D9F6F1"))
    ax.text(11.025, 1.45, f"Supercapacitor\npeak {metrics['Supercapacitor current peak']:.2f} A ≤ 18 A",
            ha="center", va="center", fontsize=9.5, weight="bold", color="#115E59")
    ax.add_patch(FancyArrowPatch((11.03, 3.18), (11.03, 1.92), arrowstyle="<|-|>", mutation_scale=14,
                                 linewidth=1.5, color="#0F766E"))
    ax.text(7.0, 5.92, "Measured signal flow in the Simscape protection model", ha="center", va="center",
            fontsize=16, weight="bold", color="#0F172A")
    plt.tight_layout()
    fig.savefig(output, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return output


def configure_document(document):
    section = document.sections[0]
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)

    for style_name, size, color in [
        ("Title", 25, RGBColor(18, 52, 86)),
        ("Heading 1", 17, RGBColor(18, 52, 86)),
        ("Heading 2", 13, RGBColor(0, 102, 102)),
    ]:
        style = document.styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.color.rgb = color

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("PV Lightning Protection — Simscape Results")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(100, 116, 139)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Simulation study — not a certification or commercial SPD qualification")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(100, 116, 139)


def build_report():
    metrics, summary = load_data()
    create_flow_diagram(metrics, summary)
    document = Document()
    configure_document(document)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("PV Simscape Model\nStep-by-Step Surge-Protection Report")
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Physical-component model • Default 10 kA, 8/20 µs indirect-lightning case")
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 102, 102)
    generated = document.add_paragraph()
    generated.alignment = WD_ALIGN_PARAGRAPH.CENTER
    generated.add_run(f"Generated from an executed Simscape run on {date.today().isoformat()}").italic = True

    document.add_heading("Purpose", level=1)
    document.add_paragraph(
        "This short report follows voltage and current through the model in the same order as the numbered Simulink Scopes. "
        "It explains what the PV panel supplies, where the surge is injected, how SPD1 and SPD2 divert surge current to ground, "
        "and what remains at the protected DC bus, supercapacitor, relay, inverter and load."
    )

    table = document.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for cell, text in zip(table.rows[0].cells, ["Stage", "Measured result", "Meaning"]):
        shade(cell, "123456")
        set_cell_text(cell, text, bold=True, color=(255, 255, 255), size=9)
    summary_rows = [
        ("PV panel", f"{metrics['PV voltage before surge']:.2f} V, {metrics['PV current before surge']:.2f} A, {metrics['PV power delivered to MPPT/boost']:.2f} W", "Normal energy entering MPPT/boost"),
        ("Surge injection", f"{metrics['Measured injected-current peak']:.0f} A; node peak {metrics['Injection-node voltage peak']:.2f} V", "Applied at 0.600 s with an 8/20 µs shape"),
        ("SPD1", f"{metrics['SPD1 diverted-current peak']:.2f} A to ground ({summary['spd1DiversionPercent']:.2f}%)", "Primary high-current diversion"),
        ("Protected bus", f"{metrics['Protected DC-bus peak']:.2f} V", "Below 52.8 V warning and 69.6 V emergency thresholds"),
    ]
    for stage, result, meaning in summary_rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], stage, bold=True)
        set_cell_text(cells[1], result)
        set_cell_text(cells[2], meaning)

    add_picture(document, "00_system_flow.png", "Figure 1. End-to-end electrical and protection path used in this report.", 6.9)

    document.add_page_break()
    add_step_heading(document, 1, "PV panel output and MPPT/boost input")
    document.add_paragraph(
        f"Before the lightning event, the physical Solar Cell array supplies an average {metrics['PV voltage before surge']:.2f} V and "
        f"{metrics['PV current before surge']:.2f} A. Their product is approximately {metrics['PV power delivered to MPPT/boost']:.2f} W. "
        "This is the DC power presented to the MPPT-controlled boost stage."
    )
    add_bullet(document, "The MPPT controller changes the converter duty cycle to keep the panel near its maximum-power operating point.")
    add_bullet(document, f"The boost/DC network establishes an approximately {metrics['Protected DC bus before surge']:.2f} V protected bus before the surge.")
    add_bullet(document, "The lightning waveform is a separate source in the model; the PV panel does not generate the 10 kA surge.")
    add_picture(document, "01_pv_output.png", "Figure 2. Scope 01 quantities: PV voltage, current and calculated power before surge injection.")

    document.add_page_break()
    add_step_heading(document, 2, "Indirect-lightning surge injection")
    document.add_paragraph(
        f"At t = {summary['eventTime_s']:.3f} s, the controlled physical source injects an 8/20 µs line-to-ground current pulse. "
        f"The measured peak is {metrics['Measured injected-current peak']:.2f} A, matching the configured {metrics['Configured lightning-current peak']:.0f} A. "
        f"The injection node reaches {metrics['Injection-node voltage peak']:.2f} V before the coordinated protection stages reduce the disturbance."
    )
    add_bullet(document, "8 µs means the waveform reaches its peak 8 µs after the event begins.")
    add_bullet(document, "20 µs means it has decayed to 50% of peak at 20 µs.")
    add_bullet(document, "Scope 02 shows both the configured command and the current/voltage measured in the physical network.")
    add_picture(document, "02_lightning_injection.png", "Figure 3. Scope 02 evidence: configured and measured surge current, plus injection-node voltage.")

    document.add_page_break()
    add_step_heading(document, 3, "SPD1 primary diversion to ground")
    document.add_paragraph(
        f"SPD1 is the upstream MOV stage. Its model conduction knee is {metrics['SPD1 conduction threshold']:.0f} V. "
        f"As the surge rises, the nonlinear MOV becomes strongly conductive and diverts a peak {metrics['SPD1 diverted-current peak']:.2f} A to ground. "
        f"This equals {summary['spd1DiversionPercent']:.2f}% of the injected peak, leaving a peak residual of {metrics['Residual-current peak after SPD1']:.2f} A for the coordinated second stage."
    )
    add_bullet(document, "The threshold is a nonlinear conduction knee, not an ideal switch that holds exactly 62 V.")
    add_bullet(document, f"The measured SPD1 terminal peak is {metrics['SPD1 terminal-voltage peak']:.2f} V because lead/source inductance creates microsecond overshoot while current changes extremely quickly.")
    add_bullet(document, "Peak currents occur at different instants, so SPD1, SPD2 and residual peak values should not be subtracted as if they were simultaneous DC values.")
    add_picture(document, "03_spd_current_diversion.png", "Figure 4. Scopes 03–04: primary SPD1 diversion and the smaller coordinated currents at SPD2.")

    document.add_page_break()
    add_step_heading(document, 4, "SPD2 coordination and protected output voltage")
    document.add_paragraph(
        f"SPD2 is placed downstream and has a lower {metrics['SPD2 conduction threshold']:.0f} V conduction knee. "
        f"It diverts a peak {metrics['SPD2 diverted-current peak']:.2f} A to ground. Its terminal voltage peaks at "
        f"{metrics['SPD2 terminal-voltage peak']:.2f} V, and the peak residual-current measurement after SPD2 is "
        f"{metrics['Residual-current peak after SPD2']:.2f} A."
    )
    document.add_paragraph(
        f"Most importantly, the protected DC bus peaks at only {metrics['Protected DC-bus peak']:.2f} V. This is "
        f"{summary['busMargin_V']:.2f} V below the {metrics['Emergency trip threshold']:.1f} V emergency relay threshold and also below the 52.8 V warning threshold. "
        "Therefore the relay remains closed in this default short surge case."
    )
    add_bullet(document, "SPD1 handles the high-current front; SPD2 reduces the coordinated residual close to the protected equipment.")
    add_bullet(document, "The DC-link capacitance and physical impedances also influence the protected-bus voltage; the SPDs are not ideal voltage sources.")
    add_picture(document, "04_voltage_protection.png", "Figure 5. Scope 07 comparison from the high-voltage injection node to the protected DC bus.")

    document.add_page_break()
    add_step_heading(document, 5, "Supercapacitor, relay, inverter and load")
    document.add_paragraph(
        f"The supercapacitor is a slower energy buffer, not a 10 kA lightning-current path. Its maximum current over the executed run is "
        f"{metrics['Supercapacitor current peak']:.2f} A, below its {metrics['Supercapacitor current limit']:.0f} A limit. "
        "The microsecond surge is primarily handled by the MOVs and DC-link capacitance."
    )
    document.add_paragraph(
        f"Before the event, the averaged inverter supplies approximately {metrics['AC output voltage RMS before surge']:.2f} V RMS and "
        f"{metrics['AC output current RMS before surge']:.2f} A RMS to the modeled load. Because the protected bus remains below the warning and emergency thresholds, "
        "the relay command and physical contact stay at logic 1 (closed)."
    )
    add_picture(document, "05_downstream_output.png", "Figure 6. Scopes 05 and 08: supercapacitor behavior and inverter/load AC output.")

    document.add_heading("Where to see each result in Simulink", level=2)
    scope_table = document.add_table(rows=1, cols=2)
    scope_table.style = "Table Grid"
    scope_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, text in zip(scope_table.rows[0].cells, ["Scope", "What it shows"]):
        shade(cell, "006666")
        set_cell_text(cell, text, bold=True, color=(255, 255, 255))
    scope_rows = [
        ("01", "PV voltage, PV current, PV power and MPPT duty"),
        ("02", "Configured/measured lightning current and injection voltage"),
        ("03", "Injected current, SPD1 current to ground and residual after SPD1"),
        ("04", "SPD2 current, residual after SPD2 and protected-bus voltage"),
        ("05", "Supercapacitor voltage, current, power and energy"),
        ("06", "Warning/emergency thresholds, relay command and physical state"),
        ("07", "Voltage comparison: injection, SPD1, SPD2 and protected output"),
        ("08", "Protected bus, relay state and AC load voltage/current"),
    ]
    for scope, description in scope_rows:
        cells = scope_table.add_row().cells
        set_cell_text(cells[0], scope, bold=True)
        set_cell_text(cells[1], description)

    document.add_heading("Conclusion", level=2)
    document.add_paragraph(
        "The executed physical-component simulation shows the intended protection sequence: the PV/MPPT stage supplies normal DC power; "
        "the separate 10 kA indirect-lightning source creates a fast high-voltage disturbance; SPD1 sends nearly all peak surge current to ground; "
        "SPD2 handles the coordinated residual; and the protected bus remains below the warning and emergency relay thresholds while the downstream system continues operating."
    )
    note = document.add_paragraph()
    run = note.add_run("Engineering limitation: ")
    run.bold = True
    note.add_run(
        "These are model results using documented component assumptions. They are not certified surge-test results, insulation-coordination approval, or a substitute for manufacturer SPD data and laboratory validation."
    )

    document.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_report())
